import os
import re
import PyPDF2
# Для анализа структуры PDF и извлечения текста
from pdfminer.high_level import extract_pages, extract_text
from pdfminer.layout import LTTextContainer, LTChar, LTRect, LTFigure
# Для извлечения текста из таблиц в PDF
import pdfplumber
from handler.handler_def import text_extraction, crop_image, convert_to_images, image_to_text, extract_table, table_converter
from docx import Document

from handler.pulling_out import summarize_text


class HandlerPDF:
    def __init__(self):
        self.__text_per_page = {}
        self.__page_text = None
        self.__line_format = None
        self.__text_from_images = None
        self.__text_from_tables = None
        self.__page_content = None

    def extrations_from_pdf(self, pdf_path: str):
        # print(type(pdf))
        pdfFileObj = open(pdf_path, 'rb')
        pdfReaded = PyPDF2.PdfReader(pdfFileObj)

        for pagenum, page in enumerate(extract_pages(pdf_path)):

            # Инициализируем переменные, необходимые для извлечения текста со страницы
            pageObj = pdfReaded.pages[pagenum]
            self.__page_text = []
            self.__line_format = []
            self.__text_from_images = []
            self.__text_from_tables = []
            self.__page_content = []
            # Инициализируем количество исследованных таблиц
            table_num = 0
            first_element = True
            table_extraction_flag = False
            # Открываем файл pdf
            pdf = pdfplumber.open(pdf_path)
            # Находим исследуемую страницу
            page_tables = pdf.pages[pagenum]
            # Находим количество таблиц на странице
            tables = page_tables.find_tables()

            # Находим все элементы
            page_elements = [(element.y1, element) for element in page._objs]
            # Сортируем все элементы по порядку нахождения на странице
            page_elements.sort(key=lambda a: a[0], reverse=True)

            # Находим элементы, составляющие страницу
            for i, component in enumerate(page_elements):
                # Извлекаем положение верхнего края элемента в PDF
                pos = component[0]
                # Извлекаем элемент структуры страницы
                element = component[1]

                # Проверяем, является ли элемент текстовым
                if isinstance(element, LTTextContainer):
                    # Проверяем, находится ли текст в таблице
                    if table_extraction_flag == False:
                        # Используем функцию извлечения текста и формата для каждого текстового элемента
                        (line_text, format_per_line) = text_extraction(element)
                        # Добавляем текст каждой строки к тексту страницы
                        self.__page_text.append(line_text)
                        # Добавляем формат каждой строки, содержащей текст
                        self.__line_format.append(format_per_line)
                        self.__page_content.append(line_text)
                    else:
                        # Пропускаем текст, находящийся в таблице
                        pass

                # Проверяем элементы на наличие изображений
                if isinstance(element, LTFigure):
                    # Вырезаем изображение из PDF
                    crop_image(element, pageObj)
                    # Преобразуем обрезанный pdf в изображение
                    convert_to_images('OCRdemo/OCRPDF/handler/cropped_image.pdf')
                    # Извлекаем текст из изображения
                    image_text = image_to_text('/home/stas/PycharmProjects/OCRdemo/OCRPDF/handler/PDF_image.png')
                    self.__text_from_images.append(image_text)
                    self.__page_content.append(image_text)
                    # Добавляем условное обозначение в списки текста и формата
                    self.__page_text.append('image')
                    self.__line_format.append('image')

                # Проверяем элементы на наличие таблиц
                if isinstance(element, LTRect):
                    # Если первый прямоугольный элемент
                    if first_element == True and (table_num + 1) <= len(tables):
                        # Находим ограничивающий прямоугольник таблицы
                        lower_side = page.bbox[3] - tables[table_num].bbox[3]
                        upper_side = element.y1
                        # Извлекаем информацию из таблицы
                        table = extract_table(pdf_path, pagenum, table_num)
                        # Преобразуем информацию таблицы в формат структурированной строки
                        table_string = table_converter(table)
                        # Добавляем строку таблицы в список
                        self.__text_from_tables.append(table_string)
                        self.__page_content.append(table_string)
                        # Устанавливаем флаг True, чтобы избежать повторения содержимого
                        table_extraction_flag = True
                        # Преобразуем в другой элемент
                        first_element = False
                        # Добавляем условное обозначение в списки текста и формата
                        self.__page_text.append('table')
                        self.__line_format.append('table')

                    # Проверяем, извлекли ли мы уже таблицы из этой страницы
                    if element.y0 >= lower_side and element.y1 <= upper_side:
                        pass
                    elif not isinstance(page_elements[i][1], LTRect):
                        table_extraction_flag = False
                        first_element = True
                        table_num += 1

            # Создаём ключ для словаря
            dctkey = 'Page_' + str(pagenum)
            # Добавляем список списков как значение ключа страницы
            self.__text_per_page[dctkey] = [self.__page_text, self.__line_format, self.__text_from_images,
                                            self.__text_from_tables, self.__page_content]

            # result = ''.join(self.text_per_page['Page_0'][4])

        pdfFileObj.close()

        return self.__text_per_page


class DOCX:
    def __init__(self, name: str):
        self.__doc = Document()
        self.__name = name

    def add_paragraph(self, text: str) -> None:
        self.__doc.add_paragraph(text)

    def save(self, dir: str) -> None:
        self.__doc.save(f'media/{dir}/{self.__name}.docx')

    def doc(self):
        return self.__doc

    @property
    def name(self):
        return self.__name


class PdfToDocx:

    def __init__(self, name: str):
        self.__document = DOCX(name)
        self.__summary = DOCX(f'summary_{name}')
        self.__handler = HandlerPDF()
        self.__result = ''
        self.__key_words_list = ['Наименование заказчикa',
                                 'Наименование проекта',
                                 'Адрес(-а) расположения защищаемых объектов заказчика',
                                 'Перечень работ выполняемых',
                                 'Перечень требований по функциям проектируемой системы защиты информации',
                                 'Информация о объекте(-ах)защиты',
                                 'заказчик',
                                 'ТЕХНИЧЕСКОЕ ЗАДАНИЕ',
                                 'Срок',
                                 ]

    def init_recognition(self, pdf_path: str):
        page = self.__handler.extrations_from_pdf(pdf_path)
        for i, j in enumerate(page):
            b = ''.join(page[f'Page_{i}'][4])

            # for k in self.__key_words_list:
            #     if bool(re.search(r'\b{}\b'.format(re.escape(k)), b, re.IGNORECASE)):
            self.__result += b
        self.__document.add_paragraph(self.__result)
        self.__document.save("docx_files")
        self.__summary.add_paragraph(summarize_text(self.__result))
        self.__summary.save("summary")
        # os.remove('/home/stas/PycharmProjects/OCRdemo/OCRPDF/handler/cropped_image.pdf')
        # os.remove('/home/stas/PycharmProjects/OCRdemo/OCRPDF/handler/PDF_image.png')
        return self.__document.name, self.__summary.name
